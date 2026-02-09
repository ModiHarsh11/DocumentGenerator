from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
from django.conf import settings
from django.utils import timezone
from django.template.loader import render_to_string

from io import BytesIO
import json, os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from weasyprint import HTML, CSS

import google.generativeai as genai

from .constants import DESIGNATION_MAP

import uuid
import datetime

# ===============================
# DOCUMENT SELECTOR - REMOVED
# ===============================

# ---------------- GEMINI ----------------
genai.configure(api_key=settings.GEMINI_API_KEY)
gemini_model = genai.GenerativeModel("gemini-2.5-flash-lite")

# ---------------- LOAD JSON ----------------
BASE_DIR = settings.BASE_DIR
with open(os.path.join(BASE_DIR, "office_order.json"), encoding="utf-8") as f:
    OFFICE_ORDER = json.load(f)

with open(os.path.join(BASE_DIR, "circular.json"), encoding="utf-8") as f:
    CIRCULAR = json.load(f)

# ---------------- HELPERS ----------------
def format_date_ddmmyyyy(date_str):
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d-%m-%Y")
    except Exception:
        return date_str

# ---------------- HOME ----------------
def home(request):
    return render(request, "generator/home.html", {
        "designations": DESIGNATION_MAP.keys(),
        "people": CIRCULAR["people"]
    })

# =====================================================================
# ===================== OFFICE ORDER (UNCHANGED) ======================
# =====================================================================

def generate_body(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    prompt = request.POST.get("body_prompt", "").strip()
    lang = request.POST.get("language", "en")

    if lang == "hi":
        system_prompt = """
आप BISAG-N के लिए एक आधिकारिक कार्यालय आदेश की मुख्य सामग्री लिख रहे हैं।

नियम:
- कम से कम 2–3 वाक्यों का एक औपचारिक अनुच्छेद लिखें।
- सरकारी भाषा का प्रयोग करें।
- कोई शीर्षक, संदर्भ, दिनांक, प्रेषक या प्राप्तकर्ता न लिखें।
- बुलेट या क्रमांक का प्रयोग न करें।
- केवल सादा पाठ में उत्तर दें।
"""
    else:
        system_prompt = """
You are drafting the BODY of an official government Office Order for BISAG-N.

Rules:
- Write one formal paragraph (minimum 2–3 sentences).
- Use official government tone.
- Do not include title, reference, date, From or To.
- No bullet points or numbering.
- Plain text only.
"""

    res = gemini_model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
    return HttpResponse(res.text.strip())


def result_office_order(request):
    if request.method != "POST":
        return redirect("home")

    lang = request.POST.get("language", "en")
    raw_date = request.POST.get("date")
    date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")

    # Get reference from form, or use default
    reference = request.POST.get("reference", "").strip()
    if not reference:
        reference = (
            "बायसेग-एन/कार्यालय आदेश/2026/"
            if lang == "hi"
            else "BISAG-N/Office Order/2026/"
        )

    data = {
        "language": lang,
        "header": OFFICE_ORDER["header"][lang],
        "title": OFFICE_ORDER["title_hi"] if lang == "hi" else OFFICE_ORDER["title_en"],
        "reference": reference,
        "date": date,
        "body": request.POST.get("body", "").strip(),
        "from": DESIGNATION_MAP[request.POST.get("from_position")][lang],
        "to": [DESIGNATION_MAP[x][lang] for x in request.POST.getlist("to_recipients[]")],
    }

    request.session["doc_data"] = data
    return render(request, "generator/result_office_order.html", data)

# PDF + DOCX for office order → UNCHANGED
# (your existing download_pdf & download_docx remain exactly same)

# =====================================================================
# ========================= CIRCULAR (UPDATED) ========================
# =====================================================================

def circular_form(request):
    return render(request, "generator/circular_form.html", {
        "people": CIRCULAR["people"]
    })

# -------- GEMINI CIRCULAR BODY --------
def generate_circular_body(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=400)

    prompt = request.POST.get("body_prompt", "").strip()
    lang = request.POST.get("language", "en")

    if lang == "hi":
        system_prompt = """
आप BISAG-N के लिए एक सरकारी परिपत्र (Circular) का केवल मुख्य भाग (BODY) लिख रहे हैं।

महत्वपूर्ण नियम:
- केवल परिपत्र का मुख्य विषय-वस्तु लिखें।
- कोई विषय (Subject) न लिखें।
- कोई शीर्षक न लिखें।
- कोई संदर्भ संख्या न लिखें।
- कोई हस्ताक्षर न लिखें।
- कोई दिनांक न लिखें।
- कोई "प्रेषक" या "प्राप्तकर्ता" न लिखें।
- 1–2 औपचारिक अनुच्छेद लिखें।
- सरकारी भाषा का प्रयोग करें।
- केवल सादा पाठ में उत्तर दें।
"""
    else:
        system_prompt = """
You are drafting ONLY the BODY content of an official Government Circular for BISAG-N.

IMPORTANT Rules:
- Write ONLY the main body content of the circular.
- Do NOT include any subject line.
- Do NOT include any title or heading.
- Do NOT include reference number.
- Do NOT include signature.
- Do NOT include date.
- Do NOT include From or To sections.
- Write 1–2 formal paragraphs only.
- Official government tone.
- Plain text only.
"""

    res = gemini_model.generate_content(system_prompt + "\n\nTopic:\n" + prompt)
    return HttpResponse(res.text.strip())

# -------- CIRCULAR PREVIEW --------
def result_circular(request):
    if request.method != "POST":
        return redirect("circular_form")

    lang = request.POST.get("language")
    raw_date = request.POST.get("date")
    date = format_date_ddmmyyyy(raw_date) if raw_date else timezone.now().strftime("%d-%m-%Y")
    subject = request.POST.get("subject")
    body = request.POST.get("body")
    
    # Get from position from dropdown (like office order)
    from_position = request.POST.get("from_position")
    from_designation = DESIGNATION_MAP[from_position][lang] if from_position else ""
    
    to_ids = request.POST.getlist("to[]")

    people = CIRCULAR["people"]
    to_people = [p for p in people if str(p["id"]) in to_ids]

    # Get header based on language
    if lang == "hi":
        header = {
            "org_name": CIRCULAR["header"]["hindi"]["org_name"],
            "ministry": CIRCULAR["header"]["hindi"]["ministry"],
            "government": CIRCULAR["header"]["hindi"]["government"]
        }
    else:
        header = {
            "org_name": CIRCULAR["header"]["english"]["org_name"],
            "ministry": CIRCULAR["header"]["english"]["ministry"],
            "government": CIRCULAR["header"]["english"]["government"]
        }

    data = {
        "language": lang,
        "header": header,
        "date": date,
        "subject": subject,
        "body": body,
        "from": from_designation,
        "to_people": to_people,
    }

    request.session["circular_data"] = data
    return render(request, "generator/result_circular.html", data)

# -------- CIRCULAR PDF --------
def download_circular_pdf(request):
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    html = render_to_string("generator/pdf_circular.html", data)

    # Optimize PDF generation with font subsetting and compression
    pdf = HTML(
        string=html,
        base_url=settings.BASE_DIR
    ).write_pdf(
        optimize_images=True,
        jpeg_quality=85,
        presentational_hints=True
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Circular.pdf"'
    return response

# -------- CIRCULAR DOCX --------
def download_circular_docx(request):
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    data = request.session.get("circular_data")
    if not data:
        return HttpResponse("No circular generated", status=400)

    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add BISAG Logo
    logo_path = os.path.join(settings.BASE_DIR, "static", "generator", "bisag_logo.png")
    if os.path.exists(logo_path):
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.9))
        doc.add_paragraph()  # Add space after logo

    # Header lines
    for line in data["header"].values():
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)

    # Circular title
    lang = data.get("language", "en")
    title_text = "परिपत्र" if lang == "hi" else "Circular"
    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(16)
    
    # Date
    date_label = "दिनांक :" if lang == "hi" else "Date :"
    p = doc.add_paragraph(f"{date_label} {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Subject
    subject_label = "विषय :" if lang == "hi" else "Subject :"
    p = doc.add_paragraph(f"{subject_label} {data['subject']}")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    # Body
    p = doc.add_paragraph(data["body"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.runs[0].font.size = Pt(12)

    # From section
    p = doc.add_paragraph(data["from"])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Add some space before table
    doc.add_paragraph()
    
    # To section - Table
    to_people = data.get("to_people", [])
    if to_people:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        sr_label = "क्र." if lang == "hi" else "Sr. No."
        name_label = "नाम" if lang == "hi" else "Name"
        sign_label = "हस्ताक्षर" if lang == "hi" else "Sign"
        
        hdr_cells[0].text = sr_label
        hdr_cells[1].text = name_label
        hdr_cells[2].text = sign_label
        
        # Make header bold
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for idx, person in enumerate(to_people, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            name = person.get("name_hi") if lang == "hi" else person.get("name_en")
            row_cells[1].text = name or ""
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_cells[2].text = ""  # Empty for signature
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(1.5)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="Circular.docx"'
    return response

# ===============================
# OFFICE ORDER FORM (MISSING FIX)
# ===============================
def office_order_form(request):
    return render(request, "generator/office_order_form.html", {
        "designations": DESIGNATION_MAP.keys()
    })

# =====================================================================
# ================= OFFICE ORDER PDF & DOCX (RESTORED) =================
# =====================================================================

def download_pdf(request):
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    html = render_to_string("generator/pdf_office_order.html", data)

    pdf = HTML(
        string=html,
        base_url=settings.BASE_DIR
    ).write_pdf(
        stylesheets=[
            CSS(string="""
                @page { size: A4; margin: 2.5cm; }
                body { font-family: serif; font-size: 12pt; line-height: 1.6; }
                .center { text-align: center; }
                .bold { font-weight: bold; }
                .ref-date-row { display: table; width: 100%; margin: 20px 0; }
                .ref-left { display: table-cell; text-align: left; font-weight: bold; width: 50%; }
                .date-right { display: table-cell; text-align: right; font-weight: bold; width: 50%; }
                .title { text-align: center; font-weight: bold; text-decoration: underline; margin: 20px 0; }
                .body { text-align: justify; margin: 20px 0; }
                .from-section { text-align: right; font-weight: bold; margin: 40px 0 20px; }
                .to-section { margin-top: 20px; }
                .to-section div { margin: 5px 0; }
            """)
        ]
    )

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="Office_Order.pdf"'
    return response


def download_docx(request):
    data = request.session.get("doc_data")
    if not data:
        return HttpResponse("No office order generated", status=400)

    doc = Document()

    # Header
    for line in data["header"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

    # Reference & Date
    p = doc.add_paragraph(f"Ref: {data['reference']}")
    p.runs[0].bold = True

    p = doc.add_paragraph(f"Date: {data['date']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    # Title
    p = doc.add_paragraph(data["title"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True

    # Body
    doc.add_paragraph(data["body"])

    # From
    p = doc.add_paragraph(data["from"])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    # To
    for t in data["to"]:
        p = doc.add_paragraph(t)
        p.runs[0].bold = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="Office_Order.docx"'
        }
    )
